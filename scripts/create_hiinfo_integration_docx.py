from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT = "docs/HIINFO_SAFE_LINK_INTEGRATION_REQUEST_20260522.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    run.font.size = Pt(9.5)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(text)
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    run.bold = True
    run.font.size = Pt(15 if level == 1 else 12)
    run.font.color.rgb = RGBColor(15, 23, 42)
    return p


def add_body(doc, text, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    run.font.size = Pt(10)
    run.bold = bold
    run.font.color.rgb = RGBColor(31, 41, 55)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style=None)
        p.paragraph_format.left_indent = Cm(0.45)
        p.paragraph_format.first_line_indent = Cm(-0.25)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(f"- {item}")
        run.font.name = "Malgun Gothic"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        run.font.size = Pt(9.8)
        run.font.color.rgb = RGBColor(31, 41, 55)


def add_code_block(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F1F5F9")
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(15, 23, 42)
    doc.add_paragraph()


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, bold=True, color=(255, 255, 255))
        set_cell_shading(table.rows[0].cells[i], "1E3A8A")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    doc.add_paragraph()
    return table


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    styles = doc.styles
    styles["Normal"].font.name = "Malgun Gothic"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    styles["Normal"].font.size = Pt(10)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(6)
    r = title.add_run("SAFE-LINK 하이정보 연동 요청사항")
    r.font.name = "Malgun Gothic"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    r.font.size = Pt(20)
    r.bold = True
    r.font.color.rgb = RGBColor(15, 23, 42)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(12)
    r = sub.add_run("근로자 검증 API 연동 협의 문서 | 2026-05-22")
    r.font.name = "Malgun Gothic"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(71, 85, 105)

    add_code_block(
        doc,
        "핵심 요청: SAFE-LINK는 하이정보 DB 전체 동기화가 아니라, 최소 정보 기반의 근로자 등록 여부 검증 API를 요청합니다.",
    )

    add_heading(doc, "1. 연동 목적")
    add_body(
        doc,
        "SAFE-LINK는 건설현장 근로자의 QR/NFC 입장, 다국어 TBM, 건강 체크, 교육 확인, 전자서명, 안전문서 자동 생성을 처리하는 현장 안전관리 SaaS입니다.",
    )
    add_body(
        doc,
        "하이정보와의 연동 목적은 SAFE-LINK에서 입력된 근로자 정보를 기준으로 하이정보 DB에 등록된 근로자인지 검증하는 것입니다.",
    )
    add_body(
        doc,
        "SAFE-LINK는 하이정보 DB 전체를 가져오거나 동기화하려는 목적이 아니며, 하이정보는 등록 여부를 검증해주는 API 역할이면 충분합니다.",
    )

    add_heading(doc, "2. 권장 연동 방식")
    add_bullets(
        doc,
        [
            "근로자가 SAFE-LINK QR/NFC 화면에서 최소 정보를 입력합니다.",
            "SAFE-LINK 서버에서 입력값을 서명 또는 암호화한 뒤 하이정보 API를 호출합니다.",
            "하이정보는 DB에서 일치 여부를 검증하고 최소 결과만 응답합니다.",
            "SAFE-LINK는 검증 성공 시 자체 현장 세션을 생성하고 이후 안전관리 절차를 처리합니다.",
        ],
    )

    add_heading(doc, "3. 입력값 및 응답값")
    add_table(
        doc,
        ["구분", "필드", "설명"],
        [
            ["요청", "site_code", "하이정보 현장 코드"],
            ["요청", "name_initials", "이름 이니셜 또는 이름 일부"],
            ["요청", "phone_last4", "휴대폰 번호 뒤 4자리"],
            ["요청", "birth_date_hint", "필요 시 생년월일 일부 또는 사번 일부"],
            ["요청", "requested_at", "요청 시각"],
            ["요청", "nonce", "재전송 공격 방지를 위한 1회성 값"],
            ["요청", "signature", "HMAC-SHA256 서명값"],
            ["응답", "matched", "일치 여부"],
            ["응답", "worker_ref", "하이정보 내부 근로자 참조값"],
            ["응답", "status", "active, inactive 등 상태값"],
            ["응답", "reason_code", "실패 사유 코드"],
        ],
    )

    add_heading(doc, "4. API 예시")
    add_body(doc, "SAFE-LINK에서 하이정보로 요청하는 JSON 예시입니다.")
    add_code_block(
        doc,
        '{\n'
        '  "site_code": "GT001",\n'
        '  "name_initials": "KDH",\n'
        '  "phone_last4": "1234",\n'
        '  "birth_date_hint": "900101",\n'
        '  "requested_at": "2026-05-22T10:00:00+09:00",\n'
        '  "nonce": "random-unique-value",\n'
        '  "signature": "HMAC_SHA256_SIGNATURE"\n'
        '}',
    )
    add_body(doc, "성공 응답 예시입니다.")
    add_code_block(
        doc,
        '{\n'
        '  "matched": true,\n'
        '  "worker_ref": "HIINFO-928312",\n'
        '  "site_code": "GT001",\n'
        '  "status": "active"\n'
        '}',
    )
    add_body(doc, "실패 응답 예시입니다.")
    add_code_block(doc, '{\n  "matched": false,\n  "reason_code": "NOT_FOUND"\n}')

    add_heading(doc, "5. 보안 방식")
    add_body(doc, "API Key와 Secret은 SAFE-LINK 프론트 화면에 노출하지 않습니다.", bold=True)
    add_bullets(
        doc,
        [
            "SAFE-LINK 서버 API에서만 하이정보 API를 호출합니다.",
            "HTTPS를 필수로 사용합니다.",
            "API Key, Timestamp, Nonce, HMAC-SHA256 서명 방식을 권장합니다.",
            "Timestamp와 Nonce로 재전송 공격을 방지합니다.",
            "불필요한 개인정보는 요청 및 응답에서 제외합니다.",
        ],
    )
    add_body(doc, "권장 헤더 예시입니다.")
    add_code_block(
        doc,
        "X-API-Key: issued-api-key\n"
        "X-Timestamp: 2026-05-22T10:00:00+09:00\n"
        "X-Nonce: random-unique-value\n"
        "X-Signature: generated-hmac-signature\n"
        "Content-Type: application/json",
    )

    add_heading(doc, "6. SAFE-LINK 서버 구조")
    add_code_block(
        doc,
        "SAFE-LINK 화면\n"
        "-> SAFE-LINK 서버 API\n"
        "-> 하이정보 API\n"
        "-> SAFE-LINK DB 저장",
    )
    add_body(
        doc,
        "SAFE-LINK 서버 API는 Supabase Edge Function, Cloudflare Worker, 또는 Next.js API Route로 구성할 수 있습니다. 현재 SAFE-LINK는 Supabase를 사용 중이며 검증 결과와 세션 정보는 Supabase에 저장합니다.",
    )

    add_heading(doc, "7. SAFE-LINK 저장 기준")
    add_bullets(
        doc,
        [
            "하이정보 근로자 참조값 worker_ref",
            "현장 ID 또는 현장 코드",
            "검증 성공 여부 및 검증 시각",
            "검증 방식",
            "근로자 언어",
            "QR/NFC 세션 ID",
            "TBM 확인 기록",
            "건강 체크 기록",
            "교육 확인 기록",
            "전자서명 기록",
        ],
    )

    add_heading(doc, "8. 최초 본인 확인 및 개인정보 처리 기준")
    add_body(
        doc,
        "현장 근로자는 최초 1회 이름, 생년월일, 전화번호 뒤 4자리를 입력합니다. 현장 코드는 QR/NFC 접속 정보로 자동 식별하거나, 필요 시 선택값으로 받을 수 있습니다.",
    )
    add_body(
        doc,
        "SAFE-LINK 서버는 입력값을 하이정보 API로 전달하여 하이정보 DB에 등록된 근로자가 맞는지 여부만 검증합니다. 하이정보 응답은 일치 여부, worker_ref, 현장 코드, 상태값 등 최소 정보만 받는 것을 원칙으로 합니다.",
    )
    add_body(
        doc,
        "근로자 휴대폰 화면에는 이름과 본인 확인 완료 상태만 표시합니다. 생년월일과 전화번호 뒤 4자리는 화면에 표시하지 않습니다.",
        bold=True,
    )
    add_table(
        doc,
        ["항목", "권장 저장 방식", "근로자 화면 표시"],
        [
            ["이름", "표시용 저장 또는 마스킹 저장", "이름 또는 마스킹 이름"],
            ["생년월일", "원문 저장 지양, 해시 저장 권장", "표시하지 않음"],
            ["전화번호 뒤 4자리", "원문 저장 지양, 해시 저장 권장", "표시하지 않음"],
            ["하이정보 worker_ref", "저장", "표시하지 않음"],
            ["검증 성공 여부", "저장", "본인 확인 완료"],
            ["현장/시간/세션", "저장", "필요 시 최소 표시"],
        ],
    )
    add_body(
        doc,
        "다시 복호화할 필요가 없는 값은 암호화보다 해시 저장이 더 적합합니다. 법적 또는 운영상 원문 확인이 필요한 값만 서버 환경변수 기반 암호화 키로 암호화 저장합니다.",
    )

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_heading(doc, "9. 예상 트래픽")
    add_table(
        doc,
        ["항목", "기준"],
        [
            ["1개 현장 출력 인원", "하루 약 300명"],
            ["전체 현장", "약 25개 현장"],
            ["전체 일일 검증", "약 6,000명 내외"],
            ["월 검증 요청", "약 180,000회"],
            ["데이터 용량", "월 수 GB 이하 예상"],
        ],
    )
    add_body(
        doc,
        "보수적으로 1명당 여러 번 호출하더라도 월 수백만 회 이하 수준입니다. API 데이터 용량보다 호출 제한, 장애 대응, 개인정보 최소화, 감사 로그가 더 중요합니다.",
    )

    add_heading(doc, "10. 장애 대응")
    add_bullets(
        doc,
        [
            "하이정보 API 장애 시 관리자 승인 기반 임시 QR 입장을 검토합니다.",
            "오프라인 또는 지연 검증 상태를 저장합니다.",
            "검증 실패 및 지연 로그를 남깁니다.",
            "하이정보 API 복구 후 재검증을 수행할 수 있도록 설계합니다.",
        ],
    )

    add_heading(doc, "11. 하이정보 측 확인 요청사항")
    add_table(
        doc,
        ["번호", "확인 항목"],
        [
            ["1", "근로자 검증 API 제공 가능 여부"],
            ["2", "현장 코드, 이름, 휴대폰 뒤 4자리 등 요청 필드 기준"],
            ["3", "생년월일 일부 또는 사번 필요 여부"],
            ["4", "응답 필드 기준"],
            ["5", "API Key 및 HMAC-SHA256 가능 여부"],
            ["6", "개발 서버 URL 및 운영 서버 URL"],
            ["7", "IP 제한 필요 여부"],
            ["8", "초당/일일 호출 제한"],
            ["9", "장애 시 응답 코드"],
            ["10", "테스트용 현장 코드 및 테스트 근로자 데이터 제공 가능 여부"],
            ["11", "개인정보 최소 응답 방식 가능 여부"],
        ],
    )

    add_heading(doc, "12. 최종 요청")
    add_body(
        doc,
        "SAFE-LINK는 하이정보 DB 전체 연동이 아니라, 최소 정보 기반의 근로자 검증 API를 요청드립니다.",
        bold=True,
    )
    add_code_block(
        doc,
        "근로자 QR/NFC 접속\n"
        "-> 하이정보 등록 여부 검증\n"
        "-> SAFE-LINK에서 건강 체크\n"
        "-> 다국어 TBM/교육 확인\n"
        "-> 전자서명\n"
        "-> 안전문서 자동 생성",
    )

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
