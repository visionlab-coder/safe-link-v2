from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "generated"
SRC = OUT / "대우건설_HyperSafety_AI_공모전_공문서형_제출본_v3_가독성개선_20260611.docx"
DST = OUT / "대우건설_HyperSafety_AI_공모전_공문서형_제출본_v4_페이지압축_20260613.docx"


def set_cell_margins(cell, top=35, start=45, bottom=35, end=45):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def compact():
    doc = Document(SRC)
    for section in doc.sections:
        section.top_margin = Cm(1.15)
        section.bottom_margin = Cm(1.05)
        section.left_margin = Cm(1.25)
        section.right_margin = Cm(1.25)
        section.header_distance = Cm(0.55)
        section.footer_distance = Cm(0.55)

    for p in doc.paragraphs:
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(1.5)
        p.paragraph_format.line_spacing = 1.03

    for table in doc.tables:
        table.autofit = True
        for row in table.rows:
            row.height = None
            for cell in row.cells:
                set_cell_margins(cell)
                for p in cell.paragraphs:
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(0)
                    p.paragraph_format.line_spacing = 1.0

    doc.save(DST)
    with ZipFile(DST) as z:
        bad = z.testzip()
        if bad:
            raise RuntimeError(bad)
    print(SRC)
    print(DST)


if __name__ == "__main__":
    compact()
