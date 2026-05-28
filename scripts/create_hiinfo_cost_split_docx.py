from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT = "docs/HIINFO_SAFE_LINK_COST_SPLIT_20260522.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None, size=9.4):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.font.name = "Malgun Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(text)
    r.font.name = "Malgun Gothic"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    r.bold = True
    r.font.size = Pt(15 if level == 1 else 12)
    r.font.color.rgb = RGBColor(15, 23, 42)
    return p


def add_body(doc, text, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    r.font.name = "Malgun Gothic"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    r.font.size = Pt(10)
    r.bold = bold
    r.font.color.rgb = RGBColor(31, 41, 55)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.45)
        p.paragraph_format.first_line_indent = Cm(-0.25)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(f"- {item}")
        r.font.name = "Malgun Gothic"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        r.font.size = Pt(9.8)
        r.font.color.rgb = RGBColor(31, 41, 55)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_text(cell, header, bold=True, color=(255, 255, 255))
        set_cell_shading(cell, "1E3A8A")
        if widths:
            cell.width = Cm(widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
            if widths:
                cells[i].width = Cm(widths[i])
    doc.add_paragraph()
    return table


def add_note(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F1F5F9")
    set_cell_text(cell, text, bold=True, color=(15, 23, 42), size=9.8)
    doc.add_paragraph()


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    doc.styles["Normal"].font.name = "Malgun Gothic"
    doc.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    doc.styles["Normal"].font.size = Pt(10)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(6)
    r = title.add_run("SAFE-LINK / 하이정보 연동 데이터 및 비용 분담 산정")
    r.font.name = "Malgun Gothic"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    r.font.size = Pt(19)
    r.bold = True
    r.font.color.rgb = RGBColor(15, 23, 42)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(12)
    r = sub.add_run("일 6,000명 출력인원 기준 | 2026-05-22")
    r.font.name = "Malgun Gothic"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(71, 85, 105)

    add_note(
        doc,
        "요약: 하이정보는 등록 여부 검증 API와 DB 조회 부하만 부담하고, 서원토건 SAFE-LINK는 세션, 서명, 교육, 문서, 리포트 저장과 운영을 부담합니다.",
    )

    add_heading(doc, "1. 산정 기준")
    add_table(
        doc,
        ["항목", "기준"],
        [
            ["일 출력인원", "6,000명"],
            ["월 기준", "30일"],
            ["월 검증 인원", "180,000명"],
            ["하이정보 연동 방식", "최소 정보 기반 등록 여부 검증"],
            ["SAFE-LINK 저장 방식", "검증 결과, 세션, TBM, 건강 체크, 교육 확인, 서명, 문서 저장"],
            ["근로자 인증 방식", "Supabase Auth 계정이 아닌 worker session 방식 권장"],
        ],
        widths=[4, 11],
    )

    add_heading(doc, "2. 역할 및 부담 분리")
    add_table(
        doc,
        ["구분", "서원토건 SAFE-LINK 부담", "하이정보 부담"],
        [
            ["근로자 입력 화면", "구현 및 운영", "없음"],
            ["이름/생년월일/전화 뒤 4자리 수집", "수집 후 서버 처리", "없음"],
            ["서명/암호화/HMAC", "SAFE-LINK 서버에서 처리", "검증용 서명 확인"],
            ["하이정보 DB 조회", "API 요청 발생", "DB 조회 수행"],
            ["검증 결과 저장", "저장", "선택적 API 로그만"],
            ["TBM/교육/건강체크/서명", "전체 부담", "없음"],
            ["문서 자동 생성", "PDF/Excel/Word/HWP 생성 및 저장", "없음"],
            ["안전일지/ESG 리포트", "전체 부담", "없음"],
        ],
        widths=[3.2, 6.5, 5.3],
    )

    add_heading(doc, "3. 하이정보 부담 산정")
    add_table(
        doc,
        ["항목", "예상량"],
        [
            ["API 요청 수", "월 180,000회"],
            ["일 API 요청 수", "일 6,000회"],
            ["평균 요청", "시간당 약 250회"],
            ["출근 피크 2시간 집중 시", "초당 약 1회 미만"],
            ["보수적 피크 대비", "초당 10~30회 대응 권장"],
            ["API 송수신 데이터", "월 약 0.9~1.8GB"],
            ["API 로그 저장 시", "월 약 0.3~1GB"],
            ["DB 조회", "월 180,000회"],
        ],
        widths=[5, 10],
    )
    add_body(
        doc,
        "하이정보 입장에서는 데이터 트래픽과 DB 조회 부하는 매우 작은 편입니다. 일반적인 업무 시스템 기준으로는 운영비보다 API 개발, 보안 적용, 테스트 대응 인건비가 핵심 부담입니다.",
    )

    add_heading(doc, "4. 서원토건 SAFE-LINK 부담 산정")
    add_table(
        doc,
        ["항목", "월 예상량"],
        [
            ["검증 결과/세션/로그 DB", "약 1~2GB"],
            ["건강 체크/TBM/교육/서명 메타데이터", "DB 용량에 포함"],
            ["서명 이미지 Storage", "약 3.6~9GB"],
            ["통합 문서/리포트 Storage", "약 2~10GB"],
            ["개인별 PDF 전부 저장 시 추가", "약 27~54GB"],
            ["권장 운영 방식 총량", "월 약 10~28GB 증가"],
            ["개인별 PDF 전부 저장 방식 총량", "월 약 38~82GB 증가"],
        ],
        widths=[6, 9],
    )
    add_body(
        doc,
        "권장 방식은 개인별 문서를 전부 즉시 저장하지 않고, 필요 시 생성하거나 발급분만 저장하는 방식입니다. 이 경우 월 10~30GB 수준으로 관리 가능합니다.",
    )

    add_heading(doc, "5. Supabase 비용 기준")
    add_table(
        doc,
        ["항목", "기준"],
        [
            ["Pro 기본", "월 $25"],
            ["Pro 포함 DB Disk", "8GB 포함, 초과 $0.125/GB/월"],
            ["Pro 포함 Storage", "100GB 포함, 초과 $0.021/GB/월"],
            ["Pro 포함 Egress", "250GB 포함, 초과 $0.09/GB"],
            ["Pro 포함 Edge Function", "월 200만 회 포함, 초과 100만 회당 $2"],
            ["Compute Micro", "약 $10/월, Pro의 $10 Compute Credit으로 상쇄 가능"],
            ["Compute Small", "약 $15/월"],
            ["Compute Medium", "약 $60/월"],
        ],
        widths=[5.2, 9.8],
    )

    add_heading(doc, "6. 서원토건 월 비용 예상")
    add_body(doc, "환율은 단순 계산용으로 1달러 1,350원을 가정했습니다.")
    add_table(
        doc,
        ["운영 방식", "Supabase 구성", "예상 월 비용"],
        [
            ["POC/초기", "Pro + Micro", "약 $25, 약 34,000원"],
            ["25개 현장 안정 운영", "Pro + Small", "약 $30, 약 41,000원"],
            ["피크 여유 확보", "Pro + Medium", "약 $75, 약 101,000원"],
            ["매우 보수 운영", "Pro + Large", "약 $125, 약 169,000원"],
        ],
        widths=[4.5, 5.2, 5.3],
    )

    add_heading(doc, "7. 저장량 초과 비용 예시")
    add_table(
        doc,
        ["구분", "총량", "추가 비용"],
        [
            ["Storage", "100GB 이하", "$0"],
            ["Storage", "200GB", "초과 100GB x $0.021 = $2.10/월"],
            ["Storage", "500GB", "초과 400GB x $0.021 = $8.40/월"],
            ["Storage", "1TB", "초과 약 900GB x $0.021 = 약 $18.90/월"],
            ["DB Disk", "8GB 이하", "$0"],
            ["DB Disk", "20GB", "초과 12GB x $0.125 = $1.50/월"],
            ["DB Disk", "50GB", "초과 42GB x $0.125 = $5.25/월"],
        ],
        widths=[3.5, 4.5, 7],
    )

    add_heading(doc, "8. 중요 주의사항")
    add_body(
        doc,
        "근로자 180,000명을 Supabase Auth 사용자로 만들면 안 됩니다.",
        bold=True,
    )
    add_body(
        doc,
        "Supabase Pro는 MAU 100,000명 포함이며, 초과 시 MAU당 $0.00325가 발생합니다. 근로자 180,000명이 전부 Auth MAU로 잡히면 월 약 $260 추가 비용이 발생할 수 있습니다.",
    )
    add_bullets(
        doc,
        [
            "관리자, 본사, 현장공무만 Supabase Auth 사용자로 관리합니다.",
            "근로자는 worker_sessions, nfc_workers, qr guest session, worker_ref 방식으로 처리합니다.",
            "서명 이미지와 문서는 DB가 아니라 Storage에 저장합니다.",
            "개인별 PDF는 전부 즉시 저장하지 말고 필요 시 생성 또는 발급분만 저장하는 방식을 우선합니다.",
        ],
    )

    add_heading(doc, "9. 최종 정리")
    add_table(
        doc,
        ["회사", "월 데이터 부담", "월 비용 부담"],
        [
            ["하이정보", "API 0.9~1.8GB, 로그 0.3~1GB, DB 조회 18만 회", "운영비는 미미. API 개발/보안 적용 작업비가 핵심"],
            ["서원토건 SAFE-LINK", "DB 1~2GB 증가, Storage 10~30GB 증가 권장", "Pro + Small/Medium 기준 약 $30~75/월"],
            ["서원토건 보수 운영", "개인별 PDF 전부 저장 시 38~82GB/월 증가", "대체로 $75~125/월 범위에서 시작 가능"],
        ],
        widths=[3.2, 6.6, 5.2],
    )

    add_note(
        doc,
        "판단: 하이정보는 별도 서버 증설이 필요할 정도의 데이터량은 아니며, 서원토건은 Supabase Pro + Small로 시작하고 POC 확대 후 Medium으로 올리는 방식이 현실적입니다.",
    )

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
